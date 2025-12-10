"""
Document Processor for RAG System
Handles PDF, DOCX, TXT, and other document formats

Part of: AI Architect Portfolio Project
Project: Month 2 Week 2 - RAG Systems with LanceDB
Date: December 2025

This module provides utilities for loading and preprocessing
documents from various sources for RAG ingestion.
"""

import os
from typing import List, Dict, Optional
from pathlib import Path
import re

# Try to import document processing libraries
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  PyPDF2 not installed. PDF support disabled. Install with: pip install PyPDF2")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. DOCX support disabled. Install with: pip install python-docx")


class DocumentLoader:
    """Load documents from various file formats."""
    
    @staticmethod
    def load_txt(file_path: str) -> Dict[str, any]:
        """
        Load text file.
        
        Args:
            file_path: Path to .txt file
            
        Returns:
            Dict with text and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return {
                'text': text,
                'metadata': {
                    'source': os.path.basename(file_path),
                    'file_type': 'txt',
                    'file_path': file_path,
                    'size_bytes': os.path.getsize(file_path)
                }
            }
        except Exception as e:
            print(f"❌ Error loading TXT file {file_path}: {e}")
            return None
    
    @staticmethod
    def load_pdf(file_path: str) -> Dict[str, any]:
        """
        Load PDF file.
        
        Args:
            file_path: Path to .pdf file
            
        Returns:
            Dict with text and metadata
        """
        if not PDF_AVAILABLE:
            print("❌ PyPDF2 not installed. Cannot load PDF.")
            return None
        
        try:
            reader = PdfReader(file_path)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text_parts.append(page_text)
            
            text = '\n\n'.join(text_parts)
            
            return {
                'text': text,
                'metadata': {
                    'source': os.path.basename(file_path),
                    'file_type': 'pdf',
                    'file_path': file_path,
                    'num_pages': len(reader.pages),
                    'size_bytes': os.path.getsize(file_path)
                }
            }
        except Exception as e:
            print(f"❌ Error loading PDF file {file_path}: {e}")
            return None
    
    @staticmethod
    def load_docx(file_path: str) -> Dict[str, any]:
        """
        Load DOCX file.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Dict with text and metadata
        """
        if not DOCX_AVAILABLE:
            print("❌ python-docx not installed. Cannot load DOCX.")
            return None
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from all paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            text = '\n\n'.join(text_parts)
            
            return {
                'text': text,
                'metadata': {
                    'source': os.path.basename(file_path),
                    'file_type': 'docx',
                    'file_path': file_path,
                    'num_paragraphs': len(doc.paragraphs),
                    'size_bytes': os.path.getsize(file_path)
                }
            }
        except Exception as e:
            print(f"❌ Error loading DOCX file {file_path}: {e}")
            return None
    
    @classmethod
    def load_document(cls, file_path: str) -> Optional[Dict[str, any]]:
        """
        Auto-detect and load document based on file extension.
        
        Args:
            file_path: Path to document
            
        Returns:
            Dict with text and metadata, or None if error
        """
        ext = Path(file_path).suffix.lower()
        
        loaders = {
            '.txt': cls.load_txt,
            '.pdf': cls.load_pdf,
            '.docx': cls.load_docx,
            '.doc': cls.load_docx
        }
        
        loader = loaders.get(ext)
        if not loader:
            print(f"❌ Unsupported file type: {ext}")
            return None
        
        return loader(file_path)
    
    @classmethod
    def load_directory(
        cls,
        directory_path: str,
        file_extensions: Optional[List[str]] = None
    ) -> List[Dict[str, any]]:
        """
        Load all documents from a directory.
        
        Args:
            directory_path: Path to directory
            file_extensions: Optional list of extensions to include (e.g., ['.txt', '.pdf'])
            
        Returns:
            List of document dicts
        """
        if file_extensions is None:
            file_extensions = ['.txt', '.pdf', '.docx', '.doc']
        
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            print(f"❌ Directory not found: {directory_path}")
            return documents
        
        print(f"\n📂 Loading documents from: {directory_path}")
        print("─" * 80)
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_extensions:
                print(f"Loading: {file_path.name}")
                doc = cls.load_document(str(file_path))
                if doc:
                    documents.append(doc)
        
        print(f"\n✅ Loaded {len(documents)} documents")
        return documents


class TextCleaner:
    """Clean and preprocess text for better RAG performance."""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean text by removing excess whitespace, special characters, etc.
        
        Args:
            text: Input text
            
        Returns:
            Cleaned text
        """
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    @staticmethod
    def remove_urls(text: str) -> str:
        """Remove URLs from text."""
        return re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
    
    @staticmethod
    def remove_emails(text: str) -> str:
        """Remove email addresses from text."""
        return re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '', text)
    
    @staticmethod
    def normalize_whitespace(text: str) -> str:
        """Normalize all whitespace to single spaces."""
        return ' '.join(text.split())
    
    @classmethod
    def clean_for_rag(
        cls,
        text: str,
        remove_urls: bool = False,
        remove_emails: bool = False
    ) -> str:
        """
        Apply all cleaning steps appropriate for RAG.
        
        Args:
            text: Input text
            remove_urls: Whether to remove URLs
            remove_emails: Whether to remove emails
            
        Returns:
            Cleaned text
        """
        text = cls.clean_text(text)
        
        if remove_urls:
            text = cls.remove_urls(text)
        
        if remove_emails:
            text = cls.remove_emails(text)
        
        return text


class DocumentProcessor:
    """High-level document processing for RAG ingestion."""
    
    def __init__(self, clean_text: bool = True):
        """
        Initialize document processor.
        
        Args:
            clean_text: Whether to clean text during processing
        """
        self.loader = DocumentLoader()
        self.cleaner = TextCleaner()
        self.clean_text = clean_text
    
    def process_file(
        self,
        file_path: str,
        additional_metadata: Optional[Dict] = None
    ) -> Optional[Dict[str, any]]:
        """
        Process a single file.
        
        Args:
            file_path: Path to file
            additional_metadata: Optional metadata to add
            
        Returns:
            Processed document dict
        """
        # Load document
        doc = self.loader.load_document(file_path)
        if not doc:
            return None
        
        # Clean text if enabled
        if self.clean_text:
            doc['text'] = self.cleaner.clean_for_rag(doc['text'])
        
        # Add additional metadata
        if additional_metadata:
            doc['metadata'].update(additional_metadata)
        
        # Add word count
        doc['metadata']['word_count'] = len(doc['text'].split())
        
        return doc
    
    def process_directory(
        self,
        directory_path: str,
        file_extensions: Optional[List[str]] = None,
        additional_metadata: Optional[Dict] = None
    ) -> List[Dict[str, any]]:
        """
        Process all files in a directory.
        
        Args:
            directory_path: Path to directory
            file_extensions: Optional list of extensions to include
            additional_metadata: Optional metadata to add to all documents
            
        Returns:
            List of processed document dicts
        """
        # Load all documents
        documents = self.loader.load_directory(directory_path, file_extensions)
        
        # Process each document
        processed = []
        for doc in documents:
            if self.clean_text:
                doc['text'] = self.cleaner.clean_for_rag(doc['text'])
            
            if additional_metadata:
                doc['metadata'].update(additional_metadata)
            
            doc['metadata']['word_count'] = len(doc['text'].split())
            processed.append(doc)
        
        return processed
    
    def create_sample_documents(self) -> List[Dict[str, any]]:
        """
        Create sample documents for testing.
        
        Returns:
            List of sample document dicts
        """
        samples = [
            {
                'text': """
                Azure AI Services provide comprehensive artificial intelligence capabilities for developers.
                The platform includes Azure OpenAI Service for large language models, Cognitive Services
                for pre-built AI capabilities, and Azure Machine Learning for custom model development.
                
                Azure OpenAI Service offers access to models like GPT-4, GPT-3.5, and DALL-E 3.
                These models can be used for chatbots, content generation, code assistance, and image creation.
                The service includes enterprise-grade security, compliance, and responsible AI features.
                
                Cognitive Services provides pre-built APIs for computer vision, speech recognition,
                language understanding, and decision making. These services are easy to integrate
                and require no machine learning expertise.
                """,
                'metadata': {
                    'source': 'azure_ai_overview.txt',
                    'category': 'overview',
                    'topic': 'azure-ai'
                }
            },
            {
                'text': """
                Vector databases are specialized databases designed to store and query vector embeddings.
                Unlike traditional databases that store structured data, vector databases store
                high-dimensional vectors and enable similarity search.
                
                Popular vector databases include Pinecone, Weaviate, Chroma, and LanceDB. Each has
                different characteristics: Pinecone is fully managed, Weaviate is open-source with
                rich features, Chroma is lightweight and embedded, and LanceDB is built on Apache Arrow
                for fast performance.
                
                Vector databases are essential for RAG systems, semantic search, recommendation engines,
                and other AI applications that need to find similar items based on meaning rather than
                exact matches.
                """,
                'metadata': {
                    'source': 'vector_databases.txt',
                    'category': 'technical',
                    'topic': 'vector-db'
                }
            },
            {
                'text': """
                Retrieval-Augmented Generation (RAG) is a technique that enhances large language models
                by providing them with relevant context from a knowledge base. The RAG pipeline has
                three main stages: indexing, retrieval, and generation.
                
                During indexing, documents are split into chunks, converted to vector embeddings,
                and stored in a vector database. During retrieval, the user's query is embedded
                and used to find the most similar chunks. During generation, the retrieved chunks
                are provided as context to the LLM to generate an accurate answer.
                
                RAG solves key LLM limitations: knowledge cutoffs, hallucinations, and inability
                to access private documents. It enables LLMs to answer questions based on your
                specific documents while citing sources.
                """,
                'metadata': {
                    'source': 'rag_explained.txt',
                    'category': 'tutorial',
                    'topic': 'rag'
                }
            }
        ]
        
        # Clean text
        if self.clean_text:
            for doc in samples:
                doc['text'] = self.cleaner.clean_for_rag(doc['text'])
                doc['metadata']['word_count'] = len(doc['text'].split())
        
        return samples


# Example usage
if __name__ == "__main__":
    print("=" * 80)
    print("DOCUMENT PROCESSOR DEMO")
    print("=" * 80)
    
    processor = DocumentProcessor(clean_text=True)
    
    # Create sample documents
    print("\n📄 Creating sample documents...")
    documents = processor.create_sample_documents()
    
    print(f"\n✅ Created {len(documents)} sample documents")
    
    for i, doc in enumerate(documents, 1):
        print(f"\nDocument {i}:")
        print(f"  Source: {doc['metadata']['source']}")
        print(f"  Category: {doc['metadata']['category']}")
        print(f"  Topic: {doc['metadata']['topic']}")
        print(f"  Words: {doc['metadata']['word_count']}")
        print(f"  Preview: {doc['text'][:150]}...")
    
    print("\n" + "=" * 80)
    print("Document processing is ready for RAG ingestion!")
    print("=" * 80)
